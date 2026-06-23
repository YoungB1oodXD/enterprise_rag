# app/utils/parser.py
import re
import csv
import json as json_module
import pdfplumber
import pytesseract
import docx
import os
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

from typing import List, Dict, Any

from app.core.config import settings
from app.core.logger import get_logger

logger = get_logger(__name__)

_CHAPTER_PATTERN = re.compile(
    r'^第\s*[一二三四五六七八九十百千\d]+\s*章'
)
_SECTION_PATTERN = re.compile(
    r'^第\s*[一二三四五六七八九十百千\d]+\s*节'
)
_ARTICLE_PATTERN = re.compile(
    r'^第\s*[一二三四五六七八九十百千\d]+\s*条'
)


def extract_text_from_pdf(file_path: str) -> str:
    """
    从 PDF 中提取文本，包含扫描件 OCR 兜底。

    处理逻辑：
    - 正常 PDF：直接用 pdfplumber 提取文字
    - 扫描件（提取文字少于20字）：转图片后用 tesseract OCR
    """
    full_text = ""
    logger.info(f"开始解析PDF文件: {file_path}")
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 20:
                    full_text += f"<<PAGE:{i + 1}>>\n{page_text}\n\n"
                else:
                    logger.info(f"第 {i + 1} 页文本量少，尝试 OCR...")
                    img = page.to_image(resolution=300).original
                    try:
                        ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                        # OCR 结果同样加页码标记
                        full_text += f"<<PAGE:{i + 1}>>\n{ocr_text}\n\n"
                    except Exception as e:
                        logger.warning(f"第 {i + 1} 页 OCR 失败: {e}")
    except Exception as e:
        logger.error(f"解析 PDF 失败: {e}")
        raise
    return full_text


def extract_text_from_docx(file_path: str) -> str:
    """从 Word 文档提取文本（段落 + 表格，保持原生顺序，Markdown表格）→ RAG 工业级"""
    full_text = "<<PAGE:1>>\n\n"
    logger.info(f"开始解析 Word 文件: {file_path}")

    try:
        doc = docx.Document(file_path)

        # ==============================
        # 核心：按文档原生顺序遍历（段落 + 表格）
        # ==============================
        for child in doc.element.body:
            # 1. 段落
            if isinstance(child, CT_P):
                para = Paragraph(child, doc)
                text = para.text.strip()
                if text:
                    full_text += text + "\n\n"

            # 2. 表格（输出标准 Markdown）
            elif isinstance(child, CT_Tbl):
                table = Table(child, doc)
                full_text += "\n"

                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    full_text += "| " + " | ".join(row_data) + " |\n"

                    # 表头加分隔线
                    if row_idx == 0:
                        full_text += "|" + "|".join(["---"] * len(row.cells)) + "|\n"

                full_text += "\n"

        # ==============================
        # 增加：页眉页脚（政务必备）
        # ==============================
        full_text += "=== 页眉页脚信息 ===\n"
        for section in doc.sections:
            # 页眉
            for para in section.header.paragraphs:
                txt = para.text.strip()
                if txt:
                    full_text += f"页眉：{txt}\n"
            # 页脚
            for para in section.footer.paragraphs:
                txt = para.text.strip()
                if txt:
                    full_text += f"页脚：{txt}\n"

        full_text += "\n"

    except Exception as e:
        logger.error(f"解析 DOCX 失败: {str(e)}")
        raise

    return full_text



def chunk_text_by_headers(text: str) -> List[Dict[str, Any]]:
    """
    按政务文档的层级结构（章/节/条）进行智能分块。

    每个 chunk 包含：
    - content:    文本内容
    - breadcrumb: 面包屑路径，如 "第一章 总则 > 第三条"
    - page_number: 该 chunk 起始位置的页码
    """

    logger.info("开始政务文档结构化分块...")

    lines = text.split('\n')

    # 当前面包屑状态
    current_chapter = ""
    current_section = ""
    current_article = ""

    # 当前正在积累的 chunk
    current_content: List[str] = []
    current_page = 1
    chunk_start_page = 1

    result_chunks: List[Dict[str, Any]] = []

    def _save_current_chunk():
        content = '\n'.join(current_content).strip()
        if not content:
            return

        breadcrumb_parts = [p for p in [current_chapter, current_section, current_article] if p]
        breadcrumb = " > ".join(breadcrumb_parts) if breadcrumb_parts else "正文"

        result_chunks.append({
            "content": content,
            "breadcrumb": breadcrumb,
            "page_number": chunk_start_page,
        })

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('<<PAGE:') and stripped.endswith('>>'):
            try:
                current_page = int(stripped[7:-2])
            except ValueError:
                pass
            continue

            # 识别章级标题
        if _CHAPTER_PATTERN.match(stripped):
            _save_current_chunk()
            current_chapter = stripped
            current_section = ""
            current_article = ""
            current_content = []

            # 识别节级标题
        elif _SECTION_PATTERN.match(stripped):
            _save_current_chunk()
            current_section = stripped
            current_article = ""
            current_content = []

        # 识别条级标题（最小切分单元）
        elif _ARTICLE_PATTERN.match(stripped):
            _save_current_chunk()
            current_article = stripped
            chunk_start_page = current_page
            current_content = [line]

        # 普通正文行，追加到当前 chunk
        else:
            if not current_content:
                chunk_start_page = current_page
            current_content.append(line)

    # 循环结束后保存最后一个 chunk
    _save_current_chunk()

    final_chunks = []
    max_chunk_size = settings.rag.chunk_size     # 字符数上限，可在 config.yaml 里配置
    overlap_size = settings.rag.chunk_overlap   # 重叠字符数，防止语义在边界断裂

    for chunk in result_chunks:
        content = chunk["content"]
        if len(content) <= max_chunk_size:
            final_chunks.append(chunk)
        else:
            # 滑动窗口切分
            sub_idx = 1
            start = 0
            while start < len(content):
                end = start + max_chunk_size
                sub_content = content[start:end]
                final_chunks.append({
                    "content": sub_content,
                    # 面包屑加上子块序号，方便溯源
                    "breadcrumb": f"{chunk['breadcrumb']} (片段{sub_idx})",
                    "page_number": chunk["page_number"],
                })
                sub_idx += 1
                start = end - overlap_size  # 下一个子块从重叠位置开始

    logger.info(f"分块完成，共生成 {len(final_chunks)} 个块（原始{len(result_chunks)}个，二次切分后{len(final_chunks)}个）")
    return final_chunks


# ================================================================
# 扩展格式支持：TXT / MD / CSV / Excel / JSON / JSONL
# ================================================================
# 所有新增函数保持与现有 extract_text 相同的接口：接收 path 返回 str
# 不修改已有的 PDF/DOCX 提取逻辑
# ================================================================


def extract_text_from_txt(file_path: str) -> str:
    """
    从纯文本文件提取内容。
    优先 UTF-8，失败后尝试 GBK（中文常见编码）。
    """
    logger.info(f"开始解析 TXT 文件: {file_path}")
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            logger.info(f"TXT 解析完成（{encoding}），共 {len(content)} 字符")
            return content
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法解码 TXT 文件（尝试了 utf-8/gbk/gb2312/latin-1）: {file_path}")


def extract_text_from_markdown(file_path: str) -> str:
    """
    从 Markdown 文件提取文本。
    保留标题层级信息，用于后续按标题分块。
    """
    logger.info(f"开始解析 Markdown 文件: {file_path}")
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    # 在标题前插入 <<HEADING:level>> 标记，辅助分块
    lines = []
    for line in content.split("\n"):
        stripped = line.strip()
        heading_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            lines.append(f"<<HEADING:{level}>> {title}")
        else:
            lines.append(line)

    logger.info(f"MD 解析完成，共 {len(content)} 字符")
    return "\n".join(lines)


def extract_text_from_csv(file_path: str) -> str:
    """
    从 CSV 文件提取文本。
    每行格式为：<<RECORD:N>> 字段名1: 值1 | 字段名2: 值2 | ...
    后续 chunk_text_by_records 会按 <<RECORD:>> 标记分块。
    """
    logger.info(f"开始解析 CSV 文件: {file_path}")
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            if reader.fieldnames is None:
                raise ValueError("CSV 文件为空或格式错误")

            records = []
            for i, row in enumerate(reader, 1):
                parts = [f"{k}: {v}" for k, v in row.items() if v is not None and v.strip()]
                records.append(f"<<RECORD:{i}>> {' | '.join(parts)}")

            if not records:
                raise ValueError("CSV 文件没有数据行")

            result = "\n\n".join(records)
            logger.info(f"CSV 解析完成，共 {len(records)} 条记录")
            return result
    except UnicodeDecodeError:
        # 尝试 GBK 编码
        with open(file_path, "r", encoding="gbk") as f:
            reader = csv.DictReader(f)
            records = []
            for i, row in enumerate(reader, 1):
                parts = [f"{k}: {v}" for k, v in row.items() if v is not None and v.strip()]
                records.append(f"<<RECORD:{i}>> {' | '.join(parts)}")
            result = "\n\n".join(records)
            logger.info(f"CSV 解析完成（GBK），共 {len(records)} 条记录")
            return result


def extract_text_from_excel(file_path: str) -> str:
    """
    从 Excel (.xlsx) 文件提取文本。
    逐 sheet 解析，格式为：
      <<SHEET:N>> sheet名称
      <<RECORD:1>> 字段1: 值1 | 字段2: 值2 | ...
    """
    logger.info(f"开始解析 Excel 文件: {file_path}")
    import openpyxl

    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    except Exception as e:
        raise ValueError(f"无法打开 Excel 文件（仅支持 .xlsx 格式）: {e}")

    lines = []
    for sheet_idx, sheet_name in enumerate(wb.sheetnames, 1):
        ws = wb[sheet_name]
        lines.append(f"<<SHEET:{sheet_idx}>> {sheet_name}")

        # 读取表头（第一行）
        headers = []
        for row in ws.iter_rows(values_only=True):
            headers = [str(c) if c is not None else "" for c in row]
            break

        # 逐行读取数据
        row_count = 0
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
            if row_idx == 1:
                continue  # 跳过表头
            row_count += 1
            if headers:
                parts = []
                for j, c in enumerate(row):
                    field_name = headers[j] if j < len(headers) else f"列{j+1}"
                    val = str(c) if c is not None else ""
                    if val.strip():
                        parts.append(f"{field_name}: {val}")
                if parts:
                    lines.append(f"<<RECORD:{row_count}>> {' | '.join(parts)}")
            else:
                # 无表头时直接用值
                vals = [str(c) if c is not None else "" for c in row if c is not None]
                if vals:
                    lines.append(f"<<RECORD:{row_count}>> {' | '.join(vals)}")

        if row_count == 0:
            lines.append("（空表）")

        lines.append("")  # sheet 间空行分隔

    wb.close()
    result = "\n".join(lines).strip()
    logger.info(f"Excel 解析完成，共 {len(wb.sheetnames)} 个 sheet")
    return result


def extract_text_from_json(file_path: str) -> str:
    """
    从 JSON 文件提取文本。
    - 对象数组：每条记录为一个 <<RECORD:N>>>
    - 单对象：按一级 key 展平为一条记录
    - 嵌套结构：仅展开第一层，深层保留为 JSON 字符串
    """
    logger.info(f"开始解析 JSON 文件: {file_path}")
    with open(file_path, "r", encoding="utf-8") as f:
        data = json_module.load(f)

    def _flatten_item(item, max_depth=2) -> str:
        """将单个 JSON 对象展平为可读文本"""
        if isinstance(item, dict):
            parts = []
            for k, v in item.items():
                if isinstance(v, (dict, list)):
                    if max_depth > 1:
                        nested = _flatten_item(v, max_depth - 1)
                        if nested:
                            parts.append(f"{k}: {nested}")
                    else:
                        val_str = json_module.dumps(v, ensure_ascii=False)
                        if len(val_str) > 200:
                            val_str = val_str[:200] + "..."
                        parts.append(f"{k}: {val_str}")
                else:
                    parts.append(f"{k}: {v}")
            return " | ".join(parts)
        elif isinstance(item, list):
            return " | ".join(str(v) for v in item)
        else:
            return str(item)

    lines = []
    if isinstance(data, list):
        for i, item in enumerate(data, 1):
            text = _flatten_item(item)
            lines.append(f"<<RECORD:{i}>> {text}")
    elif isinstance(data, dict):
        lines.append(f"<<RECORD:1>> {_flatten_item(data)}")
    else:
        lines.append(f"<<RECORD:1>> {data}")

    result = "\n\n".join(lines)
    logger.info(f"JSON 解析完成，共 {len(lines)} 条记录")
    return result


def extract_text_from_jsonl(file_path: str) -> str:
    """
    从 JSONL（每行一个 JSON）文件提取文本。
    每行格式为：<<RECORD:N>> 字段1: 值1 | 字段2: 值2 | ...
    跳过程序格式和 import 语句。
    """
    logger.info(f"开始解析 JSONL 文件: {file_path}")
    lines_out = []
    record_count = 0
    with open(file_path, "r", encoding="utf-8") as f:
        for i, line in enumerate(f, 1):
            stripped = line.strip()
            if not stripped:
                continue
            try:
                data = json_module.loads(stripped)
                record_count += 1
                if isinstance(data, dict):
                    parts = [f"{k}: {v}" for k, v in data.items() if not isinstance(v, (dict, list))]
                elif isinstance(data, list):
                    parts = [str(v) for v in data]
                else:
                    parts = [str(data)]
                lines_out.append(f"<<RECORD:{record_count}>> {' | '.join(parts)}")
            except json_module.JSONDecodeError:
                continue  # 跳过非 JSON 行（注释等）

    if not lines_out:
        raise ValueError("JSONL 文件中未找到有效的 JSON 行")

    result = "\n\n".join(lines_out)
    logger.info(f"JSONL 解析完成，共 {record_count} 条记录")
    return result


# ================================================================
# 更新后的 extract_text 分发器
# ================================================================

def extract_text(file_path: str) -> str:
    """
    统一文档文本提取接口。
    根据文件后缀名自动调用对应的解析器。
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.md':
        return extract_text_from_markdown(file_path)
    elif ext == '.csv':
        return extract_text_from_csv(file_path)
    elif ext == '.xlsx':
        return extract_text_from_excel(file_path)
    elif ext in ['.json', '.jsonl']:
        return extract_text_from_jsonl(file_path) if ext == '.jsonl' else extract_text_from_json(file_path)
    else:
        raise ValueError(f"不支持的文件格式解析: {ext}")


# ================================================================
# 分块分发器 + 新分块策略
# ================================================================

def chunk_text(file_path: str, text: str) -> List[Dict[str, Any]]:
    """
    根据文件后缀选择合适的分块策略。

    - PDF / DOCX → 政务结构分块（章/节/条）
    - TXT / MD   → 段落分块（按段落 + 滑动窗）
    - CSV / Excel / JSON / JSONL → 记录分块（按 <<RECORD:N>> 标记）
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext in ['.pdf', '.docx', '.doc']:
        return chunk_text_by_headers(text)
    elif ext in ['.txt', '.md']:
        return chunk_text_by_paragraphs(text)
    elif ext in ['.csv', '.xlsx', '.json', '.jsonl']:
        return chunk_text_by_records(text)
    else:
        return chunk_text_by_paragraphs(text)


def _sliding_window_split(big_chunk: Dict[str, Any], max_size: int, overlap: int) -> List[Dict[str, Any]]:
    """
    对超大的 chunk 执行滑动窗口切分。
    复用逻辑，与 chunk_text_by_headers 的滑动窗一致。
    """
    content = big_chunk["content"]
    if len(content) <= max_size:
        return [big_chunk]

    sub_chunks = []
    sub_idx = 1
    start = 0
    while start < len(content):
        end = start + max_size
        sub_chunks.append({
            "content": content[start:end],
            "breadcrumb": f"{big_chunk['breadcrumb']} (片段{sub_idx})",
            "page_number": big_chunk["page_number"],
        })
        sub_idx += 1
        start = end - overlap

    return sub_chunks


def chunk_text_by_paragraphs(text: str) -> List[Dict[str, Any]]:
    """
    按段落分块（纯文本 / Markdown）。

    逻辑：
    1. 按双换行分割为段落
    2. 段落累积到 chunk_size 时切分
    3. 超大段落用滑动窗口二次切分
    """
    max_chunk_size = settings.rag.chunk_size
    overlap_size = settings.rag.chunk_overlap

    logger.info(f"开始段落式分块（chunk_size={max_chunk_size}, overlap={overlap_size}）...")

    # 分割段落
    paragraphs = re.split(r'\n\s*\n', text.strip())
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    if not paragraphs:
        logger.warning("无有效段落，按正文返回")
        return []

    # 按段落号分组
    raw_chunks: List[Dict[str, Any]] = []
    current_paras: List[str] = []
    current_len = 0
    start_para_idx = 1

    for i, para in enumerate(paragraphs, 1):
        if current_len + len(para) > max_chunk_size and current_paras:
            raw_chunks.append({
                "content": "\n\n".join(current_paras),
                "breadcrumb": f"段落 (第{start_para_idx}-{i-1}段)",
                "page_number": 1,
            })
            # 保留最后一段作为 overlap
            overlap_paras = []
            overlap_len = 0
            for p in reversed(current_paras):
                if overlap_len + len(p) > overlap_size and overlap_paras:
                    break
                overlap_paras.insert(0, p)
                overlap_len += len(p)
            current_paras = overlap_paras[:]
            current_len = overlap_len
            start_para_idx = i - len(overlap_paras)

        current_paras.append(para)
        current_len += len(para)

    if current_paras:
        raw_chunks.append({
            "content": "\n\n".join(current_paras),
            "breadcrumb": f"段落 (第{start_para_idx}-{start_para_idx+len(current_paras)-1}段)",
            "page_number": 1,
        })

    # 滑动窗口处理超大 chunk
    final_chunks = []
    for chunk in raw_chunks:
        final_chunks.extend(_sliding_window_split(chunk, max_chunk_size, overlap_size))

    logger.info(f"段落分块完成，共生成 {len(final_chunks)} 个块")
    return final_chunks


def chunk_text_by_records(text: str) -> List[Dict[str, Any]]:
    """
    按记录分块（CSV / Excel / JSON / JSONL）。

    逻辑：
    1. 按 <<RECORD:N>> 标记分割为独立记录
    2. 多条记录累积到 chunk_size 时切分
    3. 超大单条记录用滑动窗口二次切分
    """
    max_chunk_size = settings.rag.chunk_size
    overlap_size = settings.rag.chunk_overlap

    logger.info(f"开始记录式分块（chunk_size={max_chunk_size}, overlap={overlap_size}）...")

    # 按 <<RECORD:N>> 或 <<SHEET:N>> 分割
    record_pattern = re.compile(r'(<<(?:RECORD|SHEET):\d+>>\s*(?:.*?))(?=<<(?:RECORD|SHEET):|\Z)', re.DOTALL)
    matches = record_pattern.findall(text.strip())

    if not matches:
        # 没有结构化标记时，按段落回退
        logger.warning("未找到 RECORD 标记，回退到段落分块")
        return chunk_text_by_paragraphs(text)

    raw_chunks: List[Dict[str, Any]] = []
    current_records: List[str] = []
    current_len = 0

    for record_text in matches:
        record_text = record_text.strip()
        if not record_text:
            continue

        if current_len + len(record_text) > max_chunk_size and current_records:
            raw_chunks.append({
                "content": "\n\n".join(current_records),
                "breadcrumb": f"数据记录 (共{len(current_records)}条)",
                "page_number": 1,
            })
            current_records = []
            current_len = 0

        current_records.append(record_text)
        current_len += len(record_text)

    if current_records:
        raw_chunks.append({
            "content": "\n\n".join(current_records),
            "breadcrumb": f"数据记录 (共{len(current_records)}条)",
            "page_number": 1,
        })

    # 滑动窗口处理超大 chunk
    final_chunks = []
    for chunk in raw_chunks:
        final_chunks.extend(_sliding_window_split(chunk, max_chunk_size, overlap_size))

    logger.info(f"记录分块完成，共生成 {len(final_chunks)} 个块")
    return final_chunks